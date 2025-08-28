import io
from docx import Document
from docx.shared import Inches

def create_docx_report(report_markdown: str) -> bytes:
    """
    Takes the AI-generated markdown report and converts it into a formatted
    .docx file in memory.
    """
    document = Document()
    document.add_heading('Consolidated Security Report', 0)

    for line in report_markdown.split('\n'):
        line = line.strip()
        if line.startswith('### '):
            document.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            document.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            document.add_heading(line[2:], level=1)
        elif line.startswith('* ') or line.startswith('- '):
            document.add_paragraph(line[2:], style='List Bullet')
        elif len(line) > 0:
            document.add_paragraph(line)

    doc_stream = io.BytesIO()
    document.save(doc_stream)
    doc_stream.seek(0)
    
    return doc_stream.getvalue()